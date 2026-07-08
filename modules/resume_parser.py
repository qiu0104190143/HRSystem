"""简历解析模块 — AI大模型解析 + 规则回退"""
import os
import json
import re


def parse_resume(filepath, file_type):
    """
    解析简历文件
    优先使用AI（DeepSeek）解析，未配置API Key时回退到规则解析
    """
    # 1. 提取纯文本
    content_text = _extract_text(filepath, file_type)
    if not content_text:
        return _empty_result()

    # 2. 尝试AI解析
    ai_result = _ai_parse(content_text)
    if ai_result:
        ai_result['content_text'] = content_text
        ai_result['file_type'] = file_type
        return ai_result

    # 3. 回退到规则解析
    info = _extract_info_regex(content_text)
    info['content_text'] = content_text
    info['file_type'] = file_type
    info['parsed_json'] = json.dumps({
        'basic': {'name': info.get('name',''),'gender': info.get('gender',''),'age': info.get('age',0),'phone': info.get('phone',''),'email': info.get('email','')},
        'education': {'level': info.get('education',''),'school': info.get('school',''),'major': info.get('major','')},
        'work': {'years': info.get('experience_years',0),'current_company': info.get('current_company',''),'current_position': info.get('current_position','')},
        'skills': [s.strip() for s in info.get('skills','').replace('，',',').split(',') if s.strip()],
        'raw_text': content_text,
        'parse_method': 'regex',
    }, ensure_ascii=False)
    return info


def _ai_parse(text):
    """使用DeepSeek API解析简历"""
    try:
        from modules.models import get_setting
        api_key = get_setting('ai_api_key', '')
        api_url = get_setting('ai_api_url', 'https://api.deepseek.com/v1/chat/completions')
        model = get_setting('ai_model', 'deepseek-chat')

        if not api_key:
            return None  # API Key未配置，回退规则解析
    except Exception:
        return None

    prompt = f"""你是一个专业的简历解析专家。请从以下简历文本中提取结构化信息，严格返回JSON格式（不要任何多余文字）：

{{
  "name": "姓名",
  "gender": "男或女",
  "age": 数字,
  "phone": "手机号",
  "email": "邮箱",
  "education": "最高学历(高中/中专/大专/本科/硕士/博士/博士后)",
  "school": "毕业院校全称",
  "major": "专业名称",
  "skills": ["技能1", "技能2"],
  "experience_years": 工作年限数字,
  "current_company": "最近一家公司全称",
  "current_position": "最近一个职位",
  "summary": "一句话概括候选人亮点"
}}

规则：
- 无法确定的字段设为空字符串""或0
- skills必须是数组，尽量提取所有提到的技术技能
- education必须是以下之一：高中/中专/大专/本科/硕士/博士/博士后
- experience_years是数字，如无法确定则填0
- 只返回JSON，不要任何解释文字

简历内容：
---
{text}
---"""

    try:
        import requests
        headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json',
        }
        data = {
            'model': model,
            'messages': [
                {'role': 'system', 'content': '你是一个简历解析工具，只返回JSON，不返回其他内容。'},
                {'role': 'user', 'content': prompt},
            ],
            'temperature': 0.1,
            'max_tokens': 1000,
        }
        resp = requests.post(api_url, headers=headers, json=data, timeout=30)
        if resp.status_code != 200:
            return None

        content = resp.json()['choices'][0]['message']['content'].strip()
        # 提取JSON（处理AI可能包裹在```json```中的情况）
        if content.startswith('```'):
            content = re.sub(r'^```\w*\n?', '', content)
            content = re.sub(r'\n?```$', '', content)

        parsed = json.loads(content)

        # 整理为标准格式
        result = {
            'name': str(parsed.get('name', '') or ''),
            'gender': str(parsed.get('gender', '') or ''),
            'age': int(parsed.get('age', 0) or 0),
            'phone': str(parsed.get('phone', '') or ''),
            'email': str(parsed.get('email', '') or ''),
            'education': str(parsed.get('education', '') or ''),
            'school': str(parsed.get('school', '') or ''),
            'major': str(parsed.get('major', '') or ''),
            'skills': '，'.join(parsed.get('skills', []) if isinstance(parsed.get('skills'), list) else []),
            'experience_years': int(parsed.get('experience_years', 0) or 0),
            'current_company': str(parsed.get('current_company', '') or ''),
            'current_position': str(parsed.get('current_position', '') or ''),
            'parsed_json': json.dumps({
                'basic': {
                    'name': parsed.get('name', ''),
                    'gender': parsed.get('gender', ''),
                    'age': parsed.get('age', 0),
                    'phone': parsed.get('phone', ''),
                    'email': parsed.get('email', ''),
                },
                'education': {
                    'level': parsed.get('education', ''),
                    'school': parsed.get('school', ''),
                    'major': parsed.get('major', ''),
                },
                'work': {
                    'years': parsed.get('experience_years', 0),
                    'current_company': parsed.get('current_company', ''),
                    'current_position': parsed.get('current_position', ''),
                },
                'skills': parsed.get('skills', []) if isinstance(parsed.get('skills'), list) else [],
                'summary': str(parsed.get('summary', '') or ''),
                'raw_text': text,
                'parse_method': 'ai',
                'ai_model': model,
            }, ensure_ascii=False),
        }
        return result

    except Exception:
        return None


# ===== 文本提取（保留） =====

def _extract_text(filepath, file_type):
    """从文件提取纯文本"""
    if file_type == 'txt':
        return _parse_txt(filepath)
    elif file_type == 'docx':
        return _parse_docx(filepath)
    elif file_type == 'pdf':
        return _parse_pdf(filepath)
    return ''


def _parse_txt(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(filepath, 'r', encoding='gbk') as f:
                return f.read()
        except Exception:
            return ''


def _parse_docx(filepath):
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return '\n'.join(paragraphs)
    except Exception:
        return ''


def _parse_pdf(filepath):
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return '\n'.join(text_parts)
    except Exception:
        pass
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(filepath)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return '\n'.join(text_parts)
    except Exception:
        return ''


# ===== 规则解析（回退方案，保留原有逻辑） =====

def _extract_info_regex(text):
    """正则规则提取结构化信息（AI不可用时的回退方案）"""
    info = {
        'name': '', 'gender': '', 'age': 0, 'phone': '', 'email': '',
        'education': '', 'school': '', 'major': '', 'skills': '',
        'experience_years': 0, 'current_company': '', 'current_position': '',
    }
    lines = text.split('\n')

    # 姓名（放宽匹配：找前8行中2-8字符、不含关键词的行）
    for line in lines[:8]:
        line = line.strip()
        if line and 2 <= len(line) <= 8:
            if not any(kw in line for kw in ['简历','个人','联系','电话','邮箱','求职','应聘','姓名','性别','年龄','生日','出生','地址','民族','籍贯','政治']):
                # 中英文名均可
                if re.match(r'^[\w一-鿿·.\s]+$', line):
                    info['name'] = line
                    break

    # 邮箱
    em = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    if em: info['email'] = em.group()

    # 电话
    ph = re.search(r'1[3-9]\d{9}', text)
    if ph: info['phone'] = ph.group()

    # 性别
    if '男' in text[:200]: info['gender'] = '男'
    elif '女' in text[:200]: info['gender'] = '女'

    # 年龄
    for p in [r'年龄[：:]\s*(\d{1,2})', r'(\d{1,2})\s*岁']:
        m = re.search(p, text)
        if m and 18 <= int(m.group(1)) <= 70:
            info['age'] = int(m.group(1))
            break

    # 学历（取最高学历）
    edu_found = ''
    for edu in ['博士后','博士','硕士','研究生','本科','学士','大专','中专','高中']:
        if edu in text[:800]:
            edu_found = '硕士' if edu == '研究生' else ('本科' if edu == '学士' else edu)
            break
    info['education'] = edu_found

    # 学校（匹配完整校名含大学/学院后缀）
    for p in [r'(?:毕业院校|学校|院校)[：:]\s*([^\n]{2,30})', r'([^\n]{2,20}(?:大学|学院))']:
        m = re.search(p, text[:800])
        if m:
            s = m.group(1).strip()
            if len(s) >= 2 and not any(k in s for k in ['联系','电话','邮箱']):
                info['school'] = s
                break

    # 专业
    m = re.search(r'(?:专业)[：:]\s*([^\n]{2,20})', text[:500])
    if m: info['major'] = m.group(1).strip()

    # 技能
    info['skills'] = _extract_skills(text)

    # 经验
    for p in [r'(\d{1,2})\s*年(?:工作)?经验', r'工作经验[：:]\s*(\d{1,2})\s*年', r'工作年限[：:]\s*(\d{1,2})']:
        m = re.search(p, text)
        if m:
            info['experience_years'] = int(m.group(1))
            break

    # 公司
    m = re.search(r'(?:公司|企业)[：:]\s*([^\n]{2,30})', text[:800])
    if m: info['current_company'] = m.group(1).strip()

    # 职位
    m = re.search(r'(?:职位|岗位|职务)[：:]\s*([^\n]{2,20})', text[:800])
    if m: info['current_position'] = m.group(1).strip()

    return info


SKILL_KEYWORDS = [
    'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'Go', 'Rust',
    'PHP', 'Ruby', 'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB', 'Shell',
    'HTML', 'CSS', 'React', 'Vue', 'Angular', 'jQuery', 'Bootstrap',
    'Node.js', 'Webpack', 'Vite', 'Next.js', 'Nuxt',
    'Spring', 'Django', 'Flask', 'Express', 'FastAPI', 'MyBatis', 'Hibernate',
    '.NET', 'Laravel', 'Rails', 'Gin',
    'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Oracle', 'SQL Server',
    'SQLite', 'Elasticsearch', 'DynamoDB', 'HBase',
    'AWS', 'Azure', '阿里云', 'Docker', 'Kubernetes', 'Jenkins', 'Git',
    'CI/CD', 'Nginx', 'Linux', 'Tomcat', 'Ansible', 'Terraform',
    'Hadoop', 'Spark', 'Flink', 'Kafka', 'TensorFlow', 'PyTorch',
    '机器学习', '深度学习', 'NLP', '数据分析', '数据挖掘',
    '项目管理', '团队管理', '产品设计', 'UI/UX', 'Figma', 'Sketch',
    'Photoshop', 'Office', 'Excel', 'PPT', '沟通能力', '英语', '日语',
    '财务分析', '会计', '审计', '税务', '招聘', '培训', '绩效管理',
    '薪酬福利', '人力资源', 'HR', '财务', '出纳',
]


def _extract_skills(text):
    found = []
    tl = text.lower()
    for s in SKILL_KEYWORDS:
        if s.lower() in tl:
            found.append(s)
    return '，'.join(found)


def _empty_result():
    return {
        'name': '', 'gender': '', 'age': 0, 'phone': '', 'email': '',
        'education': '', 'school': '', 'major': '', 'skills': '',
        'experience_years': 0, 'current_company': '', 'current_position': '',
        'content_text': '', 'file_type': '', 'parsed_json': '',
    }
